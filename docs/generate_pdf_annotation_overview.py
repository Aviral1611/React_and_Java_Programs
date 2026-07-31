from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "PDF_Annotation_Version_Control_Technical_Overview.docx"
DIAGRAM = ROOT / "docs" / "pdf_annotation_architecture.png"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF7EF"
GREEN = "217A4B"
PALE_GOLD = "FFF7E0"
GOLD = "8A6500"
WHITE = "FFFFFF"
BORDER = "C9D2DC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def add_body(doc, text, bold_lead=None, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=1, after=1, line=1.08)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=accent, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    set_keep_with_next(p)
    return p


def add_numbering_definition(doc, num_id, abstract_id, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=8, line=1.167)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    header._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0, line=1.05)
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=NAVY, bold=True)

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.08)
            run = p.add_run(value)
            set_run_font(run, size=9.3, color=INK)
    set_table_geometry(table, widths_dxa)
    return table


def add_code_line(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4, line=1.0)
    p.paragraph_format.left_indent = Inches(0.22)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9.3, color=DARK_BLUE)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, before, after, size, color in (
        (1, 16, 8, 16, BLUE),
        (2, 12, 6, 13, BLUE),
        (3, 8, 4, 12, DARK_BLUE),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run("DOCUMENT VERSION CONTROL  |  PDF ANNOTATION ARCHITECTURE")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    add_numbering_definition(doc, num_id=41, abstract_id=41, bullet=False)
    add_numbering_definition(doc, num_id=42, abstract_id=42, bullet=True)


def load_font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_centered_text(draw, box, text, font, fill, max_chars=24):
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, width=max_chars) or [""])
    line_h = font.size + 7 if hasattr(font, "size") else 22
    total_h = line_h * len(lines)
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h


def arrow(draw, start, end, color=BLUE, width=7):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 12), (ex - 18 * direction, ey + 12)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 12, ey - 18 * direction), (ex + 12, ey - 18 * direction)]
    draw.polygon(points, fill=color)


def build_diagram():
    width, height = 1500, 710
    image = Image.new("RGB", (width, height), "FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = load_font(36, bold=True)
    box_title = load_font(26, bold=True)
    body_font = load_font(21)
    small_font = load_font(18, bold=True)

    draw.text((60, 35), "Runtime architecture and storage boundaries", font=title_font, fill=f"#{NAVY}")

    frontend = (70, 130, 500, 315)
    backend = (535, 130, 965, 315)
    database = (1010, 95, 1430, 260)
    files = (1010, 305, 1430, 485)
    note = (315, 565, 1185, 660)

    for box, fill, border in (
        (frontend, f"#{LIGHT_BLUE}", f"#{BLUE}"),
        (backend, "#EEF1F5", f"#{NAVY}"),
        (database, f"#{PALE_GREEN}", f"#{GREEN}"),
        (files, f"#{PALE_GOLD}", f"#{GOLD}"),
        (note, "#F8FAFC", f"#{BORDER}"),
    ):
        draw.rounded_rectangle(box, radius=20, fill=fill, outline=border, width=4)

    draw_centered_text(draw, (90, 145, 480, 205), "React frontend", box_title, f"#{NAVY}")
    draw_centered_text(
        draw,
        (105, 210, 465, 300),
        "Dashboard  |  PDF.js viewer\nAnnotation overlay  |  History",
        body_font,
        f"#{INK}",
        max_chars=32,
    )

    draw_centered_text(draw, (555, 145, 945, 205), "Java HttpServer", box_title, f"#{NAVY}")
    draw_centered_text(
        draw,
        (570, 210, 930, 300),
        "JWT + REST routes\nJDBC transaction + PDFBox",
        body_font,
        f"#{INK}",
        max_chars=30,
    )

    draw_centered_text(draw, (1030, 110, 1410, 165), "MySQL", box_title, f"#{GREEN}")
    draw_centered_text(
        draw,
        (1040, 170, 1400, 245),
        "documents\ndocument_history  |  pdf_annotations",
        body_font,
        f"#{INK}",
        max_chars=34,
    )

    draw_centered_text(draw, (1030, 320, 1410, 380), "Filesystem", box_title, f"#{GOLD}")
    draw_centered_text(
        draw,
        (1040, 385, 1400, 470),
        "backend/uploads\nCurrent PDF + history backup PDFs",
        body_font,
        f"#{INK}",
        max_chars=34,
    )

    draw_centered_text(
        draw,
        (340, 580, 1160, 645),
        "MySQL stores metadata and annotation coordinates. The filesystem stores the actual PDF bytes.",
        small_font,
        f"#{NAVY}",
        max_chars=82,
    )

    arrow(draw, (500, 222), (535, 222))
    arrow(draw, (965, 185), (1010, 185))
    arrow(draw, (965, 385), (1010, 385))
    draw.text((501, 190), "HTTP/JSON", font=small_font, fill=f"#{MUTED}")
    draw.text((968, 130), "JDBC", font=small_font, fill=f"#{GREEN}")
    draw.text((966, 330), "PDFBox / Files", font=small_font, fill=f"#{GOLD}")

    image.save(DIAGRAM, quality=95)


def add_title_block(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4, line=1.0)
    run = p.add_run("TECHNICAL OVERVIEW")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.0)
    run = p.add_run("PDF Annotation and Version Control Flow")
    set_run_font(run, size=23, color=NAVY, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=18, line=1.05)
    run = p.add_run("Document Version Control Application | React + Built-in Java HTTP Server + MySQL")
    set_run_font(run, size=13.5, color=MUTED)

    metadata = [
        ("Purpose", "Management and technical walkthrough"),
        ("Scope", "PDF upload, annotation persistence, file versioning, history download"),
        ("Implementation", "React frontend, Java HttpServer, PDFBox 3, JDBC, MySQL"),
        ("Prepared", "30 July 2026"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.0)
        lead = p.add_run(f"{label}: ")
        set_run_font(lead, size=10.5, color=INK, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=10.5, color=INK)


def build_document():
    build_diagram()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "PDF Annotation and Version Control Flow"
    doc.core_properties.subject = "Technical architecture and end-to-end data flow"
    doc.core_properties.author = "Document Version Control Project Team"
    doc.core_properties.keywords = "PDF, annotations, version control, React, Java, MySQL, PDFBox"

    add_title_block(doc)
    add_heading(doc, "Executive summary", 1)
    add_body(
        doc,
        "The application keeps each PDF in two coordinated forms. The filesystem holds the actual PDF files, while MySQL holds document metadata, version-history pointers, and annotation records. The current document row always points to one stable current PDF file. Each successful annotation save embeds the new annotations into that current PDF and also records the annotation coordinates and text in a dedicated table so the React interface can reload them.",
    )
    add_callout(
        doc,
        "Direct answer",
        "Yes, both the current PDF and all historical PDF backups are stored under the uploads folder. The annotated PDF replaces the current file at its existing path; the previous file is retained under a unique history filename. MySQL stores paths and metadata, not PDF binary data.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "The three records created around a PDF", 2)
    add_list_item(
        doc,
        "Current document: one documents row points to the latest PDF file. Its file_path remains stable while the file contents become the latest annotated version.",
        42,
        "Current document:",
    )
    add_list_item(
        doc,
        "History version: one document_history row is created for every successful annotation save and points to the backup PDF captured immediately before that save.",
        42,
        "History version:",
    )
    add_list_item(
        doc,
        "Annotation records: one pdf_annotations row is stored per comment or highlight, including page number, coordinates, dimensions or comment text, author, and timestamp.",
        42,
        "Annotation records:",
    )

    doc.add_page_break()
    add_heading(doc, "1. System architecture", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(DIAGRAM), width=Inches(6.25))
    picture._inline.docPr.set(
        "descr",
        "Architecture diagram showing the React frontend calling the Java server, which coordinates MySQL and the uploads filesystem.",
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, before=2, after=12, line=1.0)
    run = caption.add_run("Figure 1. Runtime components and storage responsibilities")
    set_run_font(run, size=9, color=MUTED, italic=True)

    add_heading(doc, "Component responsibilities", 2)
    add_table(
        doc,
        ["Component", "Primary responsibility", "What it stores or handles"],
        [
            (
                "React frontend",
                "User interaction",
                "Uploads PDFs, draws annotation overlays, sends JSON, reloads saved annotations, and downloads current/history files.",
            ),
            (
                "Java HttpServer",
                "Application coordination",
                "Authenticates requests, validates data, runs JDBC transactions, invokes PDFBox, and streams PDFs.",
            ),
            (
                "MySQL",
                "Structured metadata",
                "Document metadata, history records, annotation coordinates/text, users, timestamps, and authors.",
            ),
            (
                "uploads folder",
                "Binary file storage",
                "One current PDF per document plus uniquely named backup PDFs for prior versions.",
            ),
        ],
        [1900, 2480, 4980],
    )

    add_heading(doc, "Storage truth", 2)
    add_body(
        doc,
        "The application does not store a PDF BLOB inside MySQL. The documents.file_path column identifies the current file on disk. For a PDF history record, document_history.old_content contains the path of a previous PDF file. The field is shared with text-document history, so it contains text for a text document and a file path for a PDF document.",
    )

    doc.add_page_break()
    add_heading(doc, "2. PDF upload flow", 1)
    add_body(
        doc,
        "The upload operation establishes the stable document identity and creates the first current PDF. No history row or annotation row is required at upload time because no edit has occurred yet.",
    )
    upload_steps = [
        ("User selects a PDF.", "Dashboard.jsx validates the .pdf extension and reads the file through FileReader."),
        ("Frontend converts the file to Base64.", "The browser removes the data-URL prefix and sends filename plus fileData as JSON."),
        ("Frontend calls the upload endpoint.", "POST /api/documents/upload includes the JWT bearer token."),
        ("Backend generates a document ID.", "Java creates a UUID and names the current file <doc_id>.pdf."),
        ("Backend writes the PDF.", "The decoded bytes are saved under backend/uploads when Eclipse runs from the backend project directory."),
        ("Backend inserts the document record.", "The documents row stores title, doc_type='pdf', absolute file_path, last_updated_by, and timestamp."),
        ("Dashboard refreshes.", "GET /api/documents returns the new PDF metadata for display."),
    ]
    for title, detail in upload_steps:
        add_list_item(doc, f"{title} {detail}", 41, title)

    add_heading(doc, "State immediately after upload", 2)
    add_table(
        doc,
        ["Location", "Result"],
        [
            ("documents", "One row exists for the PDF; file_path points to backend/uploads/<doc_id>.pdf."),
            ("pdf_annotations", "No rows yet."),
            ("document_history", "No PDF history row yet."),
            ("backend/uploads", "One current PDF file exists."),
        ],
        [2150, 7210],
    )
    add_callout(
        doc,
        "Naming rule",
        "The current file uses a stable name: <doc_id>.pdf. That filename does not change after annotation saves.",
    )

    doc.add_page_break()
    add_heading(doc, "3. Annotation creation and save flow", 1)
    add_heading(doc, "Frontend behavior", 2)
    add_body(
        doc,
        "PdfAnnotator.jsx downloads the current PDF and renders its pages with PDF.js. A transparent HTML overlay sits over each page. Comments and highlight rectangles are collected in React state using PDF-space coordinates rather than being written directly by the browser.",
    )
    add_list_item(
        doc,
        "Comment: stores a UUID, type='comment', page, x, y, text, and saved=false.",
        42,
        "Comment:",
    )
    add_list_item(
        doc,
        "Highlight: stores a UUID, type='highlight', page, x, y, width, height, and saved=false.",
        42,
        "Highlight:",
    )
    add_list_item(
        doc,
        "Save request: sends only unsaved annotations to POST /api/documents/{id}/annotate.",
        42,
        "Save request:",
    )
    add_list_item(
        doc,
        "Duplicate-click protection: a React ref prevents a second save while the first request is in progress.",
        42,
        "Duplicate-click protection:",
    )

    add_heading(doc, "Backend save transaction", 2)
    save_steps = [
        ("Authenticate and parse.", "The server validates the JWT and the annotations JSON array."),
        ("Lock the document.", "SELECT ... FOR UPDATE prevents simultaneous annotation requests from overwriting each other."),
        ("Check stable annotation IDs.", "Previously saved UUIDs make a repeated successful request idempotent instead of duplicating data."),
        ("Create a temporary PDF.", "PDFBox loads the current file and adds real PDF text/highlight annotations."),
        ("Generate appearance streams.", "This makes comments and highlights visible in compatible PDF viewers."),
        ("Save and validate the temporary file.", "The output is reopened before the original PDF is touched."),
        ("Back up the current PDF.", "The previous current file is copied to a unique timestamp-and-UUID filename."),
        ("Insert history and annotations.", "The history path and annotation rows are added inside the same JDBC transaction."),
        ("Replace the current file.", "The validated temporary PDF atomically replaces <doc_id>.pdf when supported by the filesystem."),
        ("Update and commit.", "documents.last_updated_by and last_updated_at are updated, then the transaction commits."),
    ]
    for title, detail in save_steps:
        add_list_item(doc, f"{title} {detail}", 41, title)

    add_callout(
        doc,
        "Important",
        "The temporary annotated file is not retained as another permanent file. After a successful move, it becomes the current <doc_id>.pdf. Therefore the folder contains one current PDF plus one backup for every successful version change.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    doc.add_page_break()
    add_heading(doc, "4. What is stored after a successful annotation save", 1)
    add_table(
        doc,
        ["Stored item", "Example location", "Purpose", "Lifecycle"],
        [
            (
                "Current PDF",
                "uploads/<doc_id>.pdf",
                "Latest downloadable and editable PDF with all annotations embedded.",
                "Replaced on each successful save.",
            ),
            (
                "Previous PDF",
                "uploads/<doc_id>_<time>_<uuid>.pdf",
                "Exact PDF state immediately before the latest successful save.",
                "Retained for history download.",
            ),
            (
                "Document metadata",
                "documents",
                "Stable ID, title, type, current path, last editor, and last update time.",
                "One row per document.",
            ),
            (
                "History metadata",
                "document_history",
                "Links a document and editor/time to one previous PDF path.",
                "One row per successful save.",
            ),
            (
                "Annotation metadata",
                "pdf_annotations",
                "Reloads overlay positions, comments, authors, and timestamps in React.",
                "One row per annotation.",
            ),
        ],
        [1640, 2280, 3240, 2200],
    )

    add_heading(doc, "Why annotations exist in both the PDF and MySQL", 2)
    add_body(
        doc,
        "Embedding annotations into the PDF makes the downloaded file self-contained: a compatible desktop PDF viewer can display the annotations without contacting the application. Storing annotation records in MySQL serves a different purpose: the custom React viewer can quickly reconstruct its overlays, annotation list, author labels, and saved state without reverse-engineering every PDF annotation object.",
    )
    add_callout(
        doc,
        "Dual persistence",
        "PDF copy = portable document output. MySQL copy = application UI state and searchable audit metadata.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_heading(doc, "Representative database fields", 2)
    add_code_line(doc, "documents: doc_id | title | doc_type | file_path | last_updated_by | last_updated_at")
    add_code_line(doc, "document_history: history_id | doc_id | old_title | old_content(path) | changed_by | changed_at")
    add_code_line(doc, "pdf_annotations: annotation_id | doc_id | type | page | x/y | width/height | text | created_by/at")

    doc.add_page_break()
    add_heading(doc, "5. Reopening, downloading, and viewing history", 1)
    add_heading(doc, "Reopening the annotator", 2)
    add_body(
        doc,
        "When a user chooses Annotate again, the frontend starts two authenticated requests in parallel. GET /api/documents/{id}/download returns the current PDF bytes. GET /api/documents/{id}/annotations returns saved annotation rows as JSON. PDF.js renders only the base PDF page, while the React overlay renders the saved annotation records so the visual markers and annotation sidebar are recreated consistently.",
    )
    add_list_item(
        doc,
        "Saved annotations are marked saved=true and cannot be removed only from the UI, because doing so would not remove the embedded PDF object.",
        42,
    )
    add_list_item(
        doc,
        "New annotations remain saved=false and can be removed before submitting.",
        42,
    )
    add_list_item(
        doc,
        "Only new unsaved annotations are included in the next save request.",
        42,
    )

    add_heading(doc, "Current PDF download", 2)
    add_body(
        doc,
        "GET /api/documents/{id}/download looks up documents.file_path and streams the current file. Because each successful save replaces the file at that path, the user always downloads the latest annotated version.",
    )

    add_heading(doc, "History view and previous-version download", 2)
    add_body(
        doc,
        "GET /api/documents/{id}/history joins document_history with documents. For PDF entries, the server deliberately does not expose the Windows filesystem path in JSON. The React history page instead offers a Download action that calls GET /api/documents/{id}/history/{historyId}/download. The backend validates that the history entry belongs to the requested PDF and then streams the referenced backup file.",
    )
    add_callout(
        doc,
        "Current capability",
        "History versions can be reviewed and downloaded. Automatic restore of an older PDF as the current document is not implemented.",
    )

    add_heading(doc, "Relevant backend routes", 2)
    add_table(
        doc,
        ["Method and route", "Purpose"],
        [
            ("POST /api/documents/upload", "Create a new PDF document and current file."),
            ("GET /api/documents/{id}/download", "Stream the latest current PDF."),
            ("GET /api/documents/{id}/annotations", "Return saved annotation metadata."),
            ("POST /api/documents/{id}/annotate", "Embed and persist new annotations with version history."),
            ("GET /api/documents/{id}/history", "List prior text or PDF versions."),
            ("GET /api/documents/{id}/history/{historyId}/download", "Stream one previous PDF backup."),
        ],
        [4400, 4960],
    )

    doc.add_page_break()
    add_heading(doc, "6. Reliability and failure behavior", 1)
    add_body(
        doc,
        "The save order is designed to avoid the original failure mode in which history rows were committed even though PDFBox failed to write the new file.",
    )
    reliability = [
        ("Temporary-first write", "The current file is untouched until a complete annotated PDF has been saved and reopened successfully."),
        ("Database row lock", "Concurrent saves for the same document are serialized."),
        ("Transaction rollback", "History and annotation inserts are rolled back when any database or PDF step fails."),
        ("Filesystem restore", "If the current file was replaced but a later operation fails, the backup is copied back over it."),
        ("Artifact cleanup", "Failed temporary files and uncommitted backups are deleted."),
        ("Idempotent retry", "Stable annotation UUIDs prevent a repeated completed request from embedding the same annotations again."),
        ("Readable runtime errors", "PDF library linkage failures return a useful HTTP error and are logged in the Eclipse console."),
    ]
    for title, detail in reliability:
        add_list_item(doc, f"{title}: {detail}", 42, f"{title}:")

    add_heading(doc, "Operational considerations", 2)
    add_list_item(
        doc,
        "Absolute paths: documents.file_path and PDF history paths are absolute Windows paths. Moving the backend directory or database to another machine requires path migration.",
        42,
        "Absolute paths:",
    )
    add_list_item(
        doc,
        "Folder growth: every successful annotation save adds one permanent backup PDF. A future retention or archival policy may be needed for large deployments.",
        42,
        "Folder growth:",
    )
    add_list_item(
        doc,
        "Annotation history alignment: the annotation table accumulates all saved annotations, while a downloaded historical PDF contains only annotations present at that version. This is correct for viewing, but a future restore feature would also need an annotation snapshot/version link.",
        42,
        "Annotation history alignment:",
    )
    add_list_item(
        doc,
        "Shared history schema: document_history.old_content stores Markdown text for text documents and a filesystem path for PDFs. A future schema could separate text and PDF version records for stronger typing.",
        42,
        "Shared history schema:",
    )
    add_list_item(
        doc,
        "PDF coordinate edge cases: the current coordinate conversion is suitable for standard pages; rotated pages or unusual crop boxes may require viewport-transform handling.",
        42,
        "PDF coordinate edge cases:",
    )

    add_heading(doc, "7. Manager-ready explanation", 1)
    add_callout(
        doc,
        "60-second summary",
        "A PDF upload creates one database document record and one current PDF file. When a user annotates it, the browser sends annotation coordinates and text to Java. Java locks the document, builds and validates a new annotated PDF, backs up the previous current file, records the backup in version history, stores each annotation in its own table, and then replaces the current file. The current document row still points to the same path, but that path now contains the latest annotated PDF. The uploads folder therefore holds one current PDF plus all previous PDF versions.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    add_heading(doc, "Business value", 2)
    add_list_item(doc, "Users always access one clear latest version.", 42)
    add_list_item(doc, "Previous PDF states remain downloadable for audit and review.", 42)
    add_list_item(doc, "Downloaded PDFs remain self-contained because annotations are embedded.", 42)
    add_list_item(doc, "The application can reload annotations reliably because their structured data is also stored in MySQL.", 42)
    add_list_item(doc, "Atomic save and rollback behavior protects against partial updates and repeated orphan versions.", 42)

    add_heading(doc, "Implementation reference", 2)
    add_body(
        doc,
        "Frontend orchestration is implemented in PdfAnnotator.jsx, Dashboard.jsx, and DocumentHistory.jsx. Backend routing, PDFBox processing, JDBC transactions, file streaming, and history handling are implemented in backend/src/Main.java. The annotation table is defined in annotation_setup.sql.",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build_document()
    print(result)
